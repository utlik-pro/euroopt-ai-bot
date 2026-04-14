"""Импорт реальных данных от Алексея (Desktop/MVP Белхард/) в RAG.

Источники:
- Рецепты из базовой продуктовой корзины.docx (5844 параграфов)
- Короткий FAQ Еплюс_2503.docx (17 вопросов-ответов)
- Открытые Еврооопты.xlsx (798 магазинов)
"""
import sys
sys.path.insert(0, ".")

import json
import re
from pathlib import Path
from docx import Document
import openpyxl

from src.rag.engine import RAGEngine


def parse_recipes(filepath: str) -> list[dict]:
    """Парсит docx с рецептами — разбивает на отдельные рецепты."""
    doc = Document(filepath)
    recipes = []
    current_recipe = None
    current_section = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Заголовки категорий (Завтраки, Супы и т.д.)
        if para.style.name.startswith("Heading") and any(
            cat in text for cat in ["Завтраки", "Супы", "Горячие", "Гарнир", "Салат",
                                     "Выпечк", "Десерт", "Напитк", "Заготовк", "Закуск"]
        ):
            current_section = text
            continue

        # Новый рецепт (Heading 2 или жирный текст с названием блюда)
        if para.style.name.startswith("Heading") or (
            para.runs and para.runs[0].bold and len(text) < 100 and "рецепт" not in text.lower()
        ):
            if current_recipe and current_recipe.get("steps"):
                recipes.append(current_recipe)
            current_recipe = {
                "name": text,
                "category": current_section,
                "ingredients": [],
                "steps": [],
            }
            continue

        if current_recipe is None:
            continue

        # Ингредиенты (обычно с маркером или "г", "мл", "шт", "ст.л.")
        if re.search(r'\d+\s*(г|мл|шт|ст\.?л|ч\.?л|кг|л\b)', text):
            current_recipe["ingredients"].append(text)
        elif text.startswith(("•", "–", "-", "·")) or text.startswith(tuple("0123456789")):
            if "." in text[:5] or ")" in text[:5]:
                current_recipe["steps"].append(text)
            else:
                current_recipe["ingredients"].append(text)
        else:
            current_recipe["steps"].append(text)

    # Последний рецепт
    if current_recipe and current_recipe.get("steps"):
        recipes.append(current_recipe)

    return recipes


def parse_faq(filepath: str) -> list[dict]:
    """Парсит FAQ из docx (таблица: номер | вопрос+ответ)."""
    doc = Document(filepath)
    faqs = []
    for table in doc.tables:
        for row in table.rows[1:]:  # Пропускаем заголовок
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[1]:
                text = cells[1]
                # Первая строка — вопрос, остальное — ответ
                lines = text.split("\n")
                question = lines[0].strip()
                answer = "\n".join(lines[1:]).strip()
                if question and answer:
                    faqs.append({"question": question, "answer": answer})
    return faqs


def parse_stores(filepath: str) -> list[dict]:
    """Парсит xlsx с магазинами."""
    wb = openpyxl.load_workbook(filepath)
    ws = wb.active
    stores = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        name_raw = str(row[0]) if row[0] else ""
        format_type = str(row[1]) if row[1] else ""
        region = str(row[2]) if row[2] else ""

        # Извлекаем адрес из названия: "34 Магазин г.Минск,пр.Независимости,48:..."
        address_match = re.search(r'Магазин\s+(.*?)(?::\(|$)', name_raw)
        address = address_match.group(1).strip() if address_match else name_raw

        if address:
            stores.append({
                "address": address,
                "format": format_type,
                "region": region,
            })
    return stores


def main():
    rag = RAGEngine()

    # 1. Рецепты
    print("=== Загрузка рецептов ===")
    recipes = parse_recipes("data/recipes/recipes_alexey.docx")
    print(f"Найдено рецептов: {len(recipes)}")
    docs = []
    for i, r in enumerate(recipes):
        ingredients = "\n".join(r["ingredients"]) if r["ingredients"] else "не указаны"
        steps = "\n".join(r["steps"]) if r["steps"] else "не указаны"
        text = f"Рецепт: {r['name']}\nКатегория: {r['category']}\nИнгредиенты:\n{ingredients}\nПриготовление:\n{steps}"
        docs.append({
            "id": f"recipe_alexey_{i}",
            "text": text,
            "metadata": {"category": "recipes", "source": "alexey", "name": r["name"]},
        })
    rag.add_documents(docs)
    print(f"Загружено в RAG: {len(docs)}")

    # 2. FAQ Еплюс
    print("\n=== Загрузка FAQ Еплюс ===")
    faqs = parse_faq("data/faq/faq_eplus.docx")
    print(f"Найдено FAQ: {len(faqs)}")
    docs = []
    for i, f in enumerate(faqs):
        text = f"Вопрос: {f['question']}\nОтвет: {f['answer']}"
        docs.append({
            "id": f"faq_eplus_{i}",
            "text": text,
            "metadata": {"category": "faq", "source": "alexey_eplus"},
        })
    rag.add_documents(docs)
    print(f"Загружено в RAG: {len(docs)}")

    # 3. Магазины
    print("\n=== Загрузка магазинов ===")
    stores = parse_stores("data/stores/stores_alexey.xlsx")
    print(f"Найдено магазинов: {len(stores)}")
    docs = []
    for i, s in enumerate(stores):
        text = f"Магазин Евроопт: {s['address']}. Формат: {s['format']}. Регион: {s['region']}."
        docs.append({
            "id": f"store_alexey_{i}",
            "text": text,
            "metadata": {"category": "stores", "source": "alexey", "region": s["region"]},
        })
    rag.add_documents(docs)
    print(f"Загружено в RAG: {len(docs)}")

    # Итог
    total = rag.collection.count()
    print(f"\n=== ИТОГО В RAG: {total} документов ===")

    # Тест
    print("\n=== Тестовые запросы ===")
    for q in ["рецепт борща", "как получить карту Еплюс", "магазин на Независимости", "что по акции"]:
        results = rag.search(q, n_results=2)
        print(f"\nQ: {q}")
        for r in results:
            print(f"  → {r['text'][:150]}")


if __name__ == "__main__":
    main()
