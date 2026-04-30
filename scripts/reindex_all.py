"""Полная переиндексация RAG с улучшенным парсингом + новой моделью.

Источники:
- data/faq/*.json — общий FAQ + контакты
- data/faq/faq_eplus.docx — основной FAQ Еплюс (16 вопросов)
- data/promotions/*.json — товары со скидкой
- data/promotions/edostavka_raw.txt — общая инфа e-доставки
- data/recipes/sample.json + youtube.json
- data/stores/*.json + xlsx (магазины Минск + по Беларуси)

Каждый чанк сохраняем с расширенным "ключом поиска": title + alt-вопросы.
"""
import sys, json, re
from pathlib import Path
sys.path.insert(0, ".")

from docx import Document
import openpyxl
from src.rag.engine import RAGEngine

DATA = Path("data")
docs = []


def add(id_: str, text: str, category: str, **meta):
    docs.append({"id": id_, "text": text, "metadata": {"category": category, **meta}})


# === FAQ Еплюс (главный блок) ===
def load_faq_eplus():
    d = Document(DATA / "faq/faq_eplus.docx")
    n = 0
    for table in d.tables:
        for i, row in enumerate(table.rows[1:]):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2 or not cells[1]:
                continue
            text = cells[1]
            lines = text.split("\n")
            q = lines[0].strip().rstrip("?") + "?"
            a = "\n".join(lines[1:]).strip()
            if not q or not a:
                continue
            # Расширяем чанк: и вопрос, и ответ, и ключевые слова — это улучшит embedding-поиск
            full = (
                f"FAQ Еплюс — {q}\n\n"
                f"Вопрос: {q}\n\n"
                f"Ответ: {a}\n\n"
                f"Ключевые слова: Еплюс карта, как оформить карту, как получить карту, "
                f"как сделать карту, виртуальная карта, пластиковая карта, "
                f"программа лояльности, бонусы, баллы, начисление баллов, списание баллов, "
                f"купить, потратить, накопить, баланс карты"
            )
            add(f"faq_eplus_{i}", full, "faq", source="faq_eplus", topic="eplus")
            n += 1
    return n


# === FAQ general + contacts ===
def load_faq_json():
    n = 0
    for fn in ["general.json", "contacts.json", "contacts_full.json"]:
        fp = DATA / "faq" / fn
        if not fp.exists():
            continue
        items = json.load(open(fp, encoding="utf-8"))
        if isinstance(items, dict):
            items = [items]
        for i, it in enumerate(items):
            q = it.get("question") or it.get("q") or ""
            a = it.get("answer") or it.get("a") or ""
            if not q and not a:
                # contacts может быть {phone:.., email:..}
                a = " | ".join(f"{k}: {v}" for k, v in it.items())
            text = f"Контакты Евроторг" if "contact" in fn else f"Часто задаваемый вопрос"
            text += f"\n\nВопрос: {q}\nОтвет: {a}"
            add(f"faq_{fn}_{i}", text, "faq", source=fn)
            n += 1
    return n


# === Promotions (товары со скидкой) ===
def load_promotions():
    n = 0
    for fn in ["sample.json", "edostavka_crush_price.json"]:
        fp = DATA / "promotions" / fn
        if not fp.exists():
            continue
        items = json.load(open(fp, encoding="utf-8"))
        for i, it in enumerate(items):
            name = it.get("name", "")
            old = it.get("old_price", "")
            new = it.get("new_price", "")
            cat = it.get("category", "")
            disc = it.get("discount_percent", "")
            size = it.get("size", "")
            desc = it.get("description", "")
            end = it.get("end_date", "")
            text = (
                f"Акция e-доставки: {name}\n"
                f"Категория: {cat}\n"
                f"Цена: {new} руб. (было {old} руб.)"
                f"{f', скидка {disc}%' if disc else ''}"
                f"{f', {size}' if size else ''}\n"
                f"{desc}\n"
                f"Действует до: {end}\n\n"
                f"Ключевые слова: акция, скидка, промо, специальная цена, e-доставка, edostavka"
            )
            add(f"promo_{fn}_{i}", text, "promotion", source=fn, name=name, sub_category=cat)
            n += 1

    # Сырая инфа e-доставки
    raw = DATA / "promotions" / "edostavka_raw.txt"
    if raw.exists():
        content = raw.read_text(encoding="utf-8")
        # Бьём на блоки по двойным переносам, размер ~600 символов
        blocks = re.split(r"\n\n+", content)
        chunk = ""
        idx = 0
        for block in blocks:
            if len(chunk) + len(block) < 800:
                chunk += "\n\n" + block
            else:
                if chunk.strip():
                    add(f"edostavka_raw_{idx}",
                        f"E-доставка (edostavka.by) — общая информация:\n\n{chunk.strip()}",
                        "delivery", source="edostavka_raw")
                    n += 1
                    idx += 1
                chunk = block
        if chunk.strip():
            add(f"edostavka_raw_{idx}",
                f"E-доставка (edostavka.by) — общая информация:\n\n{chunk.strip()}",
                "delivery", source="edostavka_raw")
            n += 1
    return n


# === Stores Евроопт (свежий формат с координатами/расписанием) ===
def load_stores_minsk():
    n = 0
    fp = DATA / "stores/euroopt.json"
    if not fp.exists():
        return 0
    items = json.load(open(fp, encoding="utf-8"))
    for i, st in enumerate(items):
        # Новый формат после scrape (30.04.2026): brand/format/city/address/lat/lng/schedule
        # Старый формат (legacy): name/address/city/hours
        addr = st.get("address", "")
        city = st.get("city", "Минск")
        fmt = st.get("format", "")
        sched = st.get("schedule", {})
        if isinstance(sched, dict) and sched:
            hrs = (
                f"Пн-Пт {sched.get('mon','')}; "
                f"Сб {sched.get('sat','')}; "
                f"Вс {sched.get('sun','')}"
            )
        else:
            hrs = st.get("hours", "")
        text = (
            f"Магазин Евроопт в {city}\n"
            f"Формат: {fmt}\n"
            f"Адрес: {addr}\n"
            f"Часы работы: {hrs}\n\n"
            f"Ключевые слова: адрес магазина, время работы, режим, {city}, {fmt}"
        )
        add(f"store_eur_{i}", text, "store", source="euroopt.by/shops", city=city, fmt=fmt, brand="Евроопт")
        n += 1
    return n


# === Stores Беларусь (xlsx — backup на случай если scraping не сработал) ===
def load_stores_belarus():
    fp = DATA / "stores/stores_alexey.xlsx"
    # Если есть свежий all_stores.json — пропускаем xlsx, чтобы не дублировать
    if (DATA / "stores/all_stores.json").exists():
        return 0
    if not fp.exists():
        return 0
    wb = openpyxl.load_workbook(fp)
    ws = wb.active
    n = 0
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True)):
        name_raw = str(row[0]) if row[0] else ""
        fmt = str(row[1]) if row[1] else ""
        region = str(row[2]) if row[2] else ""
        m = re.search(r"Магазин\s+(.*?)(?::\(|$)", name_raw)
        addr = m.group(1).strip() if m else name_raw
        if not addr:
            continue
        text = (
            f"Магазин Евроопт\n"
            f"Адрес: {addr}\n"
            f"Формат: {fmt}\n"
            f"Регион: {region}\n\n"
            f"Ключевые слова: адрес, {region}"
        )
        add(f"store_bel_{i}", text, "store", source="stores_alexey", region=region)
        n += 1
    return n


# === Stores Хит + Грошык (из all_stores.json) ===
def load_stores_hit_groshyk():
    """Индексируем Хит и Грошык из all_stores.json (свежие данные с сайтов 30.04.2026)."""
    fp = DATA / "stores/all_stores.json"
    if not fp.exists():
        return 0
    items = json.load(open(fp, encoding="utf-8"))
    n = 0
    for i, st in enumerate(items):
        brand = st.get("brand", "")
        if brand not in ("Хит", "Грошык"):
            continue
        addr = st.get("address", "")
        city = st.get("city", "")
        region = st.get("format_group", "")
        fmt = st.get("format", "")
        text = (
            f"Магазин {brand} в {city}\n"
            f"Сеть: {brand} ({'низкие цены, дискаунтер' if brand == 'Хит' else 'жёсткий дискаунтер'})\n"
            f"Формат: {fmt}\n"
            f"Адрес: {addr}\n"
            f"Регион: {region}\n\n"
            f"Ключевые слова: адрес магазина {brand}, {city}, {region}, дискаунтер"
        )
        add(f"store_{brand.lower()}_{i}", text, "store", source=st.get("source", ""), city=city, brand=brand)
        n += 1
    return n


# === Stores автолавки ===
def load_stores_avtolavki():
    """Автолавки на сайте евроопт не показаны, но есть в нашем справочнике (55 точек)."""
    fp = DATA / "stores/all_stores.json"
    if not fp.exists():
        return 0
    items = json.load(open(fp, encoding="utf-8"))
    n = 0
    for i, st in enumerate(items):
        if st.get("format") != "Автолавка":
            continue
        addr = st.get("address", "")
        city = st.get("city", "")
        text = (
            f"Автолавка Евроопт в {city}\n"
            f"Формат: мобильный магазин (автолавка) сети Евроопт\n"
            f"Локация: {addr}\n\n"
            f"Ключевые слова: автолавка, мобильный магазин, выездная торговля, {city}, село, деревня"
        )
        add(f"store_avtolavka_{i}", text, "store", source="alexey-xlsx", city=city, fmt="Автолавка", brand="Евроопт")
        n += 1
    return n


# === Recipes — улучшенный парсинг по Heading-стилям ===
def load_recipes_docx():
    fp = DATA / "recipes/recipes_alexey.docx"
    if not fp.exists():
        return 0
    d = Document(fp)
    recipes = []
    cur = None
    for para in d.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        # Заголовок рецепта (Heading 1/2 короткий)
        is_heading = style.startswith("Heading") and len(text) < 80
        if is_heading and not text.startswith(("Шаг", "Способ", "Этап", "Вариаци")):
            if cur and cur.get("body"):
                recipes.append(cur)
            cur = {"name": text, "body": []}
            continue
        if cur is None:
            continue
        cur["body"].append(text)
    if cur and cur.get("body"):
        recipes.append(cur)

    # Фильтр: только рецепты с осмысленным названием и >= 3 строк тела
    n = 0
    for i, r in enumerate(recipes):
        name = r["name"]
        if len(r["body"]) < 3 or len(name) < 4:
            continue
        body = "\n".join(r["body"])
        text = (
            f"Рецепт: {name}\n\n{body}\n\n"
            f"Ключевые слова: рецепт, {name}, как приготовить, ингредиенты, готовить дома"
        )
        add(f"recipe_docx_{i}", text, "recipe", source="recipes_alexey", name=name)
        n += 1
    return n


def load_recipes_json():
    n = 0
    for fn in ["sample.json", "youtube.json"]:
        fp = DATA / "recipes" / fn
        if not fp.exists():
            continue
        items = json.load(open(fp, encoding="utf-8"))
        if isinstance(items, dict):
            items = [items]
        for i, r in enumerate(items):
            name = r.get("name") or r.get("title", "Рецепт")
            ingr = r.get("ingredients", [])
            steps = r.get("steps", [])
            url = r.get("url", "")
            text = f"Рецепт: {name}\n"
            if ingr:
                text += "Ингредиенты:\n" + "\n".join(f"- {i}" for i in ingr) + "\n"
            if steps:
                text += "Приготовление:\n" + "\n".join(f"{j+1}. {s}" for j, s in enumerate(steps))
            if url:
                text += f"\nВидео: {url}"
            text += f"\n\nКлючевые слова: рецепт, {name}, готовить"
            add(f"recipe_{fn}_{i}", text, "recipe", source=fn, name=name)
            n += 1
    return n


def main():
    print("=== Парсинг данных ===")
    counts = {
        "FAQ Еплюс": load_faq_eplus(),
        "FAQ JSON": load_faq_json(),
        "Promotions": load_promotions(),
        "Stores Евроопт": load_stores_minsk(),
        "Stores Беларусь (xlsx fallback)": load_stores_belarus(),
        "Stores Хит+Грошык": load_stores_hit_groshyk(),
        "Stores автолавки": load_stores_avtolavki(),
        "Recipes DOCX": load_recipes_docx(),
        "Recipes JSON": load_recipes_json(),
    }
    for k, v in counts.items():
        print(f"  {k}: {v}")
    print(f"  ИТОГО документов: {len(docs)}")

    print("\n=== Загрузка в RAG (новая коллекция) ===")
    rag = RAGEngine()
    # Сбрасываем коллекцию — старая запутана
    try:
        rag.client.delete_collection("euroopt_knowledge_v3")
        rag.client.delete_collection("euroopt_knowledge_v2")
        print("  старая коллекция удалена")
    except Exception:
        pass
    # Пересоздаём
    rag = RAGEngine()
    BATCH = 50
    for i in range(0, len(docs), BATCH):
        rag.add_documents(docs[i:i+BATCH])
    print(f"  загружено: {rag.collection.count()}")

    print("\n=== Тест поиска ===")
    for q in ["Как оформить карту Еплюс?", "Сколько баллов начисляется?",
              "Как работает e-доставка?", "Сгорают ли баллы Еплюс?",
              "Какие сейчас акции?", "Рецепт борща", "Адрес магазина в Минске"]:
        results = rag.search(q, n_results=3)
        print(f"\nQ: {q} ({len(results)})")
        for r in results[:2]:
            print(f"  score={r.get('score',0):.3f} | {(r.get('text') or '')[:120]}")


if __name__ == "__main__":
    main()
