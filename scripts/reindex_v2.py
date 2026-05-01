"""Reindex v2 — полная переиндексация с актуальными данными от Евроторга (15.04.2026).

Новое:
- Акции Евроопт 2026 (Красная цена, Выходные оптом, Цены вниз, Бонус-товары, Грильфест)
- Акции Хит Дискаунтер (Цены вниз, Жёлтые пятницы, ХИТ! ЦЕНА)
- Удача в придачу — 28 разделов FAQ + 628 товаров удачи
- Магазины с разметкой сетей Евроопт / Хит (1084 точки)
- Корректные URL: evroopt.by, groshyk.by, hitdiscount.by

Старое (сохраняем):
- FAQ Еплюс (16 вопросов)
- FAQ general + contacts
- Рецепты Алексея + YouTube
- Магазины Минск (28 с часами)
"""
import sys, json, re
from pathlib import Path
sys.path.insert(0, ".")

from docx import Document
import openpyxl
from src.rag.engine import RAGEngine

DATA = Path("data")
docs = []


def add(id_, text, category, **meta):
    docs.append({"id": id_, "text": text, "metadata": {"category": category, **meta}})


# ============ FAQ ============

def load_faq_eplus():
    d = Document(DATA / "faq/faq_eplus.docx")
    n = 0
    for table in d.tables:
        for i, row in enumerate(table.rows[1:]):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2 or not cells[1]:
                continue
            text = cells[1]
            lines = text.split("\n")
            q = lines[0].strip().rstrip("?") + "?"
            a = "\n".join(lines[1:]).strip()
            if not q or not a:
                continue
            full = (
                f"FAQ Еплюс — {q}\n\n"
                f"Вопрос: {q}\n\nОтвет: {a}\n\n"
                f"Ключевые слова: Еплюс карта, как оформить, как получить, "
                f"программа лояльности, бонусы, баллы, начисление, списание, "
                f"виртуальная карта, пластиковая карта, баланс"
            )
            add(f"faq_eplus_{i}", full, "faq", source="faq_eplus", topic="eplus")
            n += 1
    return n


def load_faq_json():
    n = 0
    for fn in ["general.json", "contacts.json", "contacts_full.json"]:
        fp = DATA / "faq" / fn
        if not fp.exists():
            continue
        items = json.load(open(fp, encoding="utf-8"))
        if isinstance(items, dict):
            items = [items]
        for i, it in enumerate(items):
            q = it.get("question") or ""
            a = it.get("answer") or ""
            # Автозамена устаревшего домена только для "грязных" источников (contacts*).
            # general.json уже вычищен вручную и может содержать осмысленные упоминания
            # "e-доставка" / "Едоставка" (в ответе-объяснении что это устарело).
            if "contact" in fn:
                a = a.replace("e-dostavka.by", "evroopt.by")
                a = a.replace("edostavka.by", "evroopt.by")
            text = "Контакты Евроторг" if "contact" in fn else "Часто задаваемый вопрос"
            text += f"\n\nВопрос: {q}\nОтвет: {a}"
            add(f"faq_{fn}_{i}", text, "faq", source=fn)
            n += 1
    return n


def load_brand_links():
    """Добавить корректные URL сетей (по требованию клиента)."""
    brands = [
        ("Евроопт", "https://evroopt.by/", "гипермаркеты продовольствия в Беларуси"),
        ("Хит Дискаунтер", "https://hitdiscount.by/", "магазины-дискаунтеры сети Евроторг"),
        ("Грошык", "https://groshyk.by/", "сеть магазинов у дома Евроторга"),
    ]
    n = 0
    for brand, url, desc in brands:
        text = (
            f"Сайт сети «{brand}»\n\n"
            f"Официальный сайт: {url}\n"
            f"Описание: {desc} сети «Евроторг».\n\n"
            f"Ключевые слова: {brand}, сайт, ссылка, официальный"
        )
        add(f"brand_link_{brand}", text, "brand", source="manual")
        n += 1
    return n


# ============ Promotions Евроопт (новые!) ============

def load_promotions_evroopt():
    fp = DATA / "promotions_new/evroopt.xlsx"
    wb = openpyxl.load_workbook(fp)
    n = 0
    # Карта синонимов для акций — улучшает BM25-матчинг парафраз
    PROMO_SYNONYMS = {
        "красная цена": "красная цена, red price, снижение цены, большие скидки, низкие цены, акция красная",
        "выходные оптом": "выходные оптом, скидки на выходных, оптом, выходные, пятница суббота воскресенье акция",
        "цены вниз": "цены вниз, снижение цен, пониженные цены, с картой Еплюс скидка, скидка по карте",
        "бонус-товар": "бонус-товары, бонусные товары, вернём бонусами, вернем бонусами, кэшбэк бонусами, возврат бонусов, до 40 процентов бонусами, процент возврата",
        "грильфест": "грильфест, шашлык, мангал, пикник, гриль, мясо для барбекю, готовим на углях",
        "х7 бонус": "Х7, Х 7, x7, x 7, семь раз, семикратные бонусы, новичкам Еплюс, для новых, регистрация Еплюс, новый участник, умножить бонусы, бонусы при регистрации, приветственные бонусы",
        "новичкам": "Х7 бонусов, новые участники, новые владельцы карты, приветственный бонус, для начинающих",
        "двойные бонус": "двойные бонусы, x2, умноженные бонусы, двойные баллы, бонусы x2, виртуальная карта двойные",
    }

    # Лист "Все акции" — описание кампаний
    ws = wb["Все акции"]
    promo_names = []
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
        name, date, desc, leaflet, page_url = row
        if not name:
            continue
        name_clean = str(name).strip()
        promo_names.append((name_clean, date))
        # Собираем синонимы
        name_lower = name_clean.lower()
        extra_kw = ""
        for key, syns in PROMO_SYNONYMS.items():
            if key in name_lower:
                extra_kw = syns
                break
        text = (
            f"Акция «Евроопт»: {name_clean}\n"
            f"Период: {date or '—'}\n"
            f"Описание: {desc or '—'}\n"
            f"Ссылка на каталог: {page_url or '—'}\n"
            f"Листовка: {leaflet or '—'}\n\n"
            f"Ключевые слова: акция Евроопт, {name_clean}, скидки, промо, листовка, акции 2026"
            + (f", {extra_kw}" if extra_kw else "")
        )
        add(f"evroopt_promo_{i}", text, "promotion",
            source="evroopt_xlsx", promo_name=name_clean[:50])
        n += 1

    # ➕ Сводный чанк «все текущие акции Евроопт» (закрывает 'какие акции сейчас?')
    promo_list_text = "\n".join(f"• {pn} — {pd}" for pn, pd in promo_names)
    summary = (
        f"Актуальные акции сети «Евроопт» (апрель 2026):\n\n"
        f"{promo_list_text}\n\n"
        f"Полный список и каталоги — на сайте https://evroopt.by/ в разделе «Акции».\n\n"
        f"Ключевые слова: все акции, какие акции, текущие акции, актуальные промо, "
        f"акции Евроопт, скидки, весенние акции, апрельские акции, что по скидкам, список акций"
    )
    add("evroopt_all_promos_summary", summary, "promotion",
        source="aggregate", promo_name="summary")
    n += 1

    # Листы с каталогами товаров
    catalogs = {
        "Каталог Красная цена": ("Красная цена", ["Код товара", "Наименование товара", "Цена без скидки", "Цена на скидки"]),
        "Каталог бонус товары": ("Бонус-товары", ["Код товара", "Наименование товара", "Цена", "Вернем рублей", "Вернем бонусов"]),
        "Каталог Грильфест": ("Грильфест", ["Код товара", "Наименование товара", "Цена"]),
        "Каталог товаров Цены Вниз": ("Цены вниз", ["Код товара", "Наименование товара", "Цена без карты", "Цена (акция)"]),
    }
    for sheet_name, (promo_name, cols) in catalogs.items():
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        current_category = None
        for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
            cells = [c for c in row if c is not None]
            if not cells:
                continue
            # Если только одна ячейка и это текст — это категория-разделитель
            if len([c for c in cells if c]) == 1 and isinstance(cells[0], str):
                current_category = cells[0]
                continue
            code = str(row[0]) if row[0] else ""
            name = str(row[1]) if len(row) > 1 and row[1] else ""
            if not name:
                continue
            price_info = ""
            if promo_name == "Красная цена":
                old, new = row[2], row[3] if len(row) > 3 else None
                price_info = f"обычная цена: {old}, по акции: {new} руб"
            elif promo_name == "Бонус-товары":
                price = row[2]
                rub_back = row[3] if len(row) > 3 else None
                bonus_back = row[4] if len(row) > 4 else None
                price_info = f"цена: {price} руб, вернём: {rub_back} руб / {bonus_back} бонусов"
            elif promo_name == "Грильфест":
                price_info = f"цена: {row[2]} руб"
            elif promo_name == "Цены вниз":
                price_info = f"без карты: {row[2]} / с акцией: {row[3]} руб"

            text = (
                f"Товар акции «{promo_name}» (Евроопт):\n"
                f"Название: {name}\n"
                f"{price_info}\n"
                f"{f'Категория: {current_category}' if current_category else ''}\n"
                f"Код: {code}\n\n"
                f"Ключевые слова: {promo_name}, акция Евроопт, скидка, товар"
            )
            add(f"evroopt_{promo_name}_{i}", text, "promotion_item",
                source="evroopt_xlsx", promo=promo_name, product=name[:60])
            n += 1
    return n


# ============ Promotions Хит Дискаунтер ============

def load_promotions_hit():
    fp = DATA / "promotions_new/hitdiscount.xlsx"
    wb = openpyxl.load_workbook(fp)
    n = 0
    # Все акции
    ws = wb["Все акции"]
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
        name, date, desc, leaflet, page_url = row
        if not name:
            continue
        text = (
            f"Акция «Хит Дискаунтер»: {str(name).strip()}\n"
            f"Период: {date or '—'}\n"
            f"Описание: {desc or '—'}\n"
            f"Ссылка на каталог: {page_url or '—'}\n\n"
            f"Ключевые слова: акция Хит, {name}, скидки, промо"
        )
        add(f"hit_promo_{i}", text, "promotion", source="hit_xlsx",
            promo_name=str(name)[:50], brand="Хит")
        n += 1

    # Каталог цены вниз
    if "Каталог цены вниз" in wb.sheetnames:
        ws = wb["Каталог цены вниз"]
        for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
            code, name, period, price = row[:4]
            if not name:
                continue
            text = (
                f"Товар акции «Цены вниз» в «Хит Дискаунтер»:\n"
                f"Название: {name}\n"
                f"Цена с картой: {price} руб\n"
                f"Период: {period}\n"
                f"Код: {code}\n\n"
                f"Ключевые слова: Хит, Цены вниз, акция, скидка"
            )
            add(f"hit_priceDown_{i}", text, "promotion_item",
                source="hit_xlsx", promo="Цены вниз", brand="Хит", product=str(name)[:60])
            n += 1
    return n


# ============ Удача в придачу ============


def load_udacha_current_tur():
    """Текущий тур (парсер главной igra.evroopt.by, обновляется по cron).

    01.05 — после теста заказчика бот выдавал тур 214 на запрос про 215.
    Теперь тянем актуальный тур из data/udacha/current.json (парсится каждый час).
    """
    fp = DATA / "udacha/current.json"
    if not fp.exists():
        return 0
    try:
        data = json.load(open(fp, encoding="utf-8"))
    except Exception:
        return 0
    if not data.get("tur_number"):
        return 0
    tur = data["tur_number"]
    prizes_summary = data.get("prizes_summary") or ""
    prizes_list = data.get("prizes_list") or []
    period = data.get("period") or ""
    draw_date = data.get("draw_date") or ""
    rules_url = data.get("rules_url") or "https://igra.evroopt.by/"

    text = (
        f"«Удача в придачу!» — ТЕКУЩИЙ ТУР {tur}\n\n"
        f"Сейчас идёт тур {tur} рекламной игры «Удача в придачу!» сети «Евроторг».\n"
    )
    if period:
        text += f"Период покупок: {period}.\n"
    if draw_date:
        text += f"Розыгрыш суперпризов: {draw_date}.\n"
    if prizes_summary:
        text += f"\n{prizes_summary}\n"
    if prizes_list:
        text += "\nПризы тура (по данным главной страницы):\n"
        for p in prizes_list[:8]:
            text += f"- {p}\n"
    text += (
        f"\nПравила игры (PDF): {rules_url}\n"
        f"Сайт игры: https://igra.evroopt.by/\n\n"
        f"ВАЖНО: денежная часть некоторых призов идёт на уплату подоходного "
        f"налога с приза и не выдаётся клиенту наличными — это публичное условие игры.\n\n"
        f"Ключевые слова: Удача в придачу, УВП, текущий тур, "
        f"тур {tur}, последний тур, актуальный тур, призы тура, "
        f"когда розыгрыш, что разыгрывается, главные призы, новый тур"
    )
    add("udacha_current_tur", text, "udacha", source="igra.evroopt.by", tur=str(tur))
    return 1


# ============ Текущая листовка Евроопт (от заказчика, snapshot) ============


def load_listovka_current():
    """Снимок текущей листовки Евроопт (data/promotions/listovka_current.json).

    Источник — xlsx от заказчика, спарсенный из официальной PDF-листовки
    evroopt.by/deals + страниц «СВАЁ» / «Родныя тавары».
    Парсится отдельным скриптом scripts/parse_listovka_xlsx.py.

    Чанки:
    1. Сводка периода + разбивка по акциям (1 чанк) — для запросов «какие акции?».
    2. Per-promo каталоги с ценами (по чанку на акцию) — для конкретных вопросов.
    3. СВАЁ ассортимент (1 чанк) — справочник кодов СВАЁ.
    4. Родныя тавары — материалы раздела (1 чанк).
    """
    fp = DATA / "promotions/listovka_current.json"
    if not fp.exists():
        return 0
    try:
        data = json.load(open(fp, encoding="utf-8"))
    except Exception:
        return 0

    products = data.get("products") or []
    meta = data.get("meta") or {}
    if not products:
        return 0

    period = meta.get("period") or "—"
    dump_date = meta.get("dump_date") or "—"
    breakdown = meta.get("promo_breakdown") or {}
    n = 0

    # 1) Сводка
    breakdown_str = "\n".join(f"  • {k}: {v} товаров" for k, v in breakdown.items())
    # Категории, по которым потом будут спрашивать («акции на молочку?»)
    # Извлекаем по корням слов из названий товаров snapshot
    name_words = " ".join(p["name"].lower() for p in products[:80])
    cat_keywords = []
    for cat, kws in [
        ("молочное (молоко, сыр, творог, кефир, сметана, масло, йогурт)",
         ["молок", "сыр", "творог", "кефир", "сметан", "масл", "йогурт"]),
        ("мясное (колбасы, свинина, говядина, курица, сосиски)",
         ["мясо", "колбас", "свинин", "говядин", "куриц", "куриное", "сосис"]),
        ("хлеб и выпечка", ["хлеб", "булк", "батон", "лепеш"]),
        ("овощи и фрукты",
         ["яблок", "томат", "помидор", "огурец", "капуст", "морков", "картоф", "банан", "апельсин"]),
        ("кондитерские изделия (торты, печенье, конфеты, шоколад)",
         ["торт", "печен", "конфет", "шокол", "пироги", "пирожн"]),
        ("напитки (сок, вода, квас, лимонад, чай, кофе)",
         ["сок", "вод", "напит", "квас", "лимонад", "чай", "кофе"]),
        ("бакалея, крупы, макароны", ["крупа", "макарон", "рис", "греч", "пшен"]),
    ]:
        if any(kw in name_words for kw in kws):
            cat_keywords.append(cat)
    cat_str = "; ".join(cat_keywords) if cat_keywords else "—"

    summary = (
        f"Текущая листовка сети «Евроопт»: {len(products)} товаров.\n"
        f"Период действия акции: {period}.\n"
        f"Дата выгрузки: {dump_date}.\n\n"
        f"Разделы листовки:\n{breakdown_str}\n\n"
        f"Категории товаров в листовке: {cat_str}.\n\n"
        f"Полная актуальная листовка с ценами и условиями: https://evroopt.by/deals/\n"
        f"Цены приведены по состоянию на начало периода и могут изменяться.\n\n"
        f"Ключевые слова: текущая листовка, актуальная листовка, действующая листовка, "
        f"актуальные акции, какие акции сейчас, что по скидкам, что в листовке, "
        f"перечень товаров со скидкой, товары со скидкой, листовка Евроопт, "
        f"акции на молочку, акции на молоко, акции на мясо, акции на хлеб, "
        f"акции на овощи, акции на фрукты, акции на бакалею, "
        f"скидки на молоко, скидки на сыр, скидки на колбасу, скидки на хлеб, "
        f"красная цена, родныя тавары, родная марка, гриль фест, грильфест, "
        f"цена вниз с Еплюс, акция Еплюс, цены с картой Еплюс, скидка по карте Еплюс, "
        f"СВАЁ, своя марка, собственная марка"
    )
    add(
        "listovka_summary",
        summary,
        "promotion",
        source="listovka_current.json",
        promo_name="summary",
    )
    n += 1

    # 2-pre) Категорийные чанки — для запросов «акции на молочку/мясо/хлеб/...»
    # Распределяем все товары по категориям через корни названий
    CAT_RULES = [
        ("молочное", ["молок", "сыр", "творог", "кефир", "сметан", "масл", "йогурт", "ряжен"],
         "молочные товары, акции на молочку, скидки на молоко, скидки на сыр, скидки на творог, кефир со скидкой, молочное со скидкой"),
        ("мясное", ["колбас", "свинин", "говядин", "куриц", "куриное", "сосис", "пельмен", "вареник", "котлет", "тушенк", "буженин", "ветчин", "бекон", "шашлык", "фарш", "карбонад", "балык"],
         "мясо, мясные товары, акции на мясо, скидки на колбасу, скидки на сосиски, мясные продукты"),
        ("хлеб", ["хлеб", "булк", "батон", "лепеш", "лаваш", "пирог", "сухар"],
         "хлеб, скидки на хлеб, выпечка, хлебобулочные"),
        ("овощи_фрукты", ["яблок", "томат", "помидор", "огурец", "капуст", "морков", "картоф", "банан", "апельсин", "лимон", "лук", "перец", "груш", "виноград", "арбуз", "дыня", "клубник", "малин", "ягод"],
         "овощи, фрукты, акции на овощи, скидки на фрукты, свежие овощи"),
        ("напитки", ["сок ", "напит", "квас", "лимонад", "вода", "пепси", "кола", "чай", "кофе", "пиво", "вино", "морс"],
         "напитки, скидки на напитки, акции на воду, сок, квас"),
        ("кондитерское", ["торт", "печен", "конфет", "шокол", "пирожн", "пряник", "вафл", "зефир", "халв", "мармелад"],
         "кондитерские, сладости, скидки на торты, конфеты, шоколад"),
    ]

    def _match_cat(name: str, kws: list[str]) -> bool:
        nm = name.lower()
        return any(kw in nm for kw in kws)

    for cat_id, kws, kw_extra in CAT_RULES:
        cat_items = [p for p in products if _match_cat(p["name"], kws)]
        if not cat_items:
            continue
        cat_items_sorted = sorted(cat_items, key=lambda x: x.get("price") or 0.0)
        head = cat_items_sorted[:25]
        lines = []
        for it in head:
            old = it.get("old_price")
            disc = it.get("discount") or ""
            line = f"• {it['name']} — {it['price']} BYN"
            if old:
                line += f" (было {old} BYN)"
            if disc:
                line += f" {disc}"
            line += f" — раздел «{it['promo']}»"
            lines.append(line)
        text = (
            f"{kw_extra}.\n\n"
            f"Категорийная подборка из текущей листовки Евроопт ({period}) — {cat_id}.\n"
            f"Всего в этой категории найдено {len(cat_items)} товаров.\n\n"
            f"Товары:\n" + "\n".join(lines) + "\n\n"
            f"Полный список с условиями — на [evroopt.by/deals](https://evroopt.by/deals/).\n\n"
            f"Ключевые слова: {kw_extra}, текущая листовка, актуальные акции, цены"
        )
        add(
            f"listovka_cat_{cat_id}",
            text,
            "promotion",
            source="listovka_current.json",
            promo_name=f"category-{cat_id}",
        )
        n += 1

    # 2) Per-promo чанки с товарами
    by_promo: dict[str, list[dict]] = {}
    for p in products:
        by_promo.setdefault(p["promo"], []).append(p)

    for promo_name, items in by_promo.items():
        # сортируем по цене для предсказуемости
        items_sorted = sorted(items, key=lambda x: x.get("price") or 0.0)
        # ограничиваем 60 товаров на чанк (RAG embedding capacity)
        head = items_sorted[:60]
        lines = []
        for it in head:
            old = it.get("old_price")
            disc = it.get("discount") or ""
            per_unit = it.get("per_unit") or ""
            until = it.get("valid_until") or period
            line = f"• {it['name']} — {it['price']} BYN"
            if old:
                line += f" (было {old} BYN)"
            if disc:
                line += f" {disc}"
            if per_unit:
                line += f"; {per_unit}"
            if until:
                line += f"; до {until}"
            note = it.get("note") or ""
            if note and "только в гипермаркетах" in note.lower():
                line += "; только в гипермаркетах"
            lines.append(line)
        body = "\n".join(lines)

        # Особое предупреждение про «Цена вниз с Еплюс»
        eplus_note = ""
        extra_kw = ""
        promo_lower = promo_name.lower()
        if "еплюс" in promo_lower:
            eplus_note = (
                "\nВНИМАНИЕ: цены этого блока действуют только при оплате с картой Еплюс. "
                "Без карты — обычная цена.\n"
            )
            extra_kw = (
                ", цены с картой Еплюс, акция по карте Еплюс, "
                "скидка по карте Еплюс, специальные цены Еплюс, "
                "что по карте Еплюс, цены вниз с Еплюс"
            )
        if "родныя" in promo_lower or "родная" in promo_lower:
            extra_kw += (
                ", Родныя тавары, родныя товары, белорусские товары, "
                "белорусское качество, родная марка"
            )
        if "сваё" in promo_lower or "сваи" in promo_lower:
            extra_kw += ", СВАЁ, своя марка, собственная марка, ассортимент СВАЁ"
        if "гриль" in promo_lower:
            extra_kw += ", грильфест, шашлык, мангал, барбекю, гриль"

        text = (
            f"Текущая листовка Евроопт — раздел «{promo_name}» ({period}).\n"
            f"Всего товаров в разделе: {len(items)}{', ниже первые ' + str(len(head)) if len(items) > len(head) else ''}.\n"
            f"{eplus_note}\n"
            f"Товары и цены:\n{body}\n\n"
            f"Полный список — на сайте https://evroopt.by/deals/.\n"
            f"Цены приведены по состоянию на начало периода и могут изменяться.\n\n"
            f"Ключевые слова: акция «{promo_name}», текущая листовка, цены, скидки, "
            f"товары акции, что в листовке, перечень товаров, какие акции, "
            f"акции сейчас, актуальные акции{extra_kw}"
        )
        # делаем безопасный id из имени промо
        slug = re.sub(r"[^a-zа-я0-9_]+", "_", promo_name.lower())[:40].strip("_")
        add(
            f"listovka_{slug}",
            text,
            "promotion",
            source="listovka_current.json",
            promo_name=promo_name[:50],
        )
        n += 1

    # 3) СВАЁ ассортимент
    svae = data.get("svae_assortment") or []
    if svae:
        lines = [f"• [{s['code']}] {s['name']} — {s.get('category','')}" for s in svae[:80]]
        text = (
            f"СВАЁ — товары собственной торговой марки сети «Евроопт».\n"
            f"Раздел сайта: https://evroopt.by/svae/\n\n"
            f"Ассортимент (фрагмент, {len(svae)} позиций):\n" + "\n".join(lines) + "\n\n"
            f"Это справочник кодов и названий; актуальные цены и наличие — "
            f"в магазинах сети и на https://evroopt.by/svae/.\n\n"
            f"Ключевые слова: СВАЁ, своё, sva, sva, своя марка, "
            f"собственная марка, евроопт сваё, частная марка, СТМ"
        )
        add(
            "listovka_svae_assortment",
            text,
            "promotion",
            source="listovka_current.json",
            promo_name="СВАЁ ассортимент",
        )
        n += 1

    # 4) Родныя тавары — новости / акции
    rod = data.get("rodnyya_news") or []
    if rod:
        lines = []
        for r in rod[:20]:
            line = f"• {r.get('date','')}: {r['title']}"
            if r.get("url"):
                line += f" — {r['url']}"
            lines.append(line)
        text = (
            f"«Родныя тавары» — белорусские товары в сети «Евроопт».\n"
            f"Раздел сайта: https://evroopt.by/narodnyya-tovary/\n\n"
            f"Материалы раздела:\n" + "\n".join(lines) + "\n\n"
            f"Ключевые слова: Родныя тавары, родныя товары, белорусские товары, "
            f"белорусское качество, родная марка, отечественные товары"
        )
        add(
            "listovka_rodnyya_news",
            text,
            "promotion",
            source="listovka_current.json",
            promo_name="Родныя тавары — новости",
        )
        n += 1

    return n


# ============ Автолавки (отдельный chunk) ============


def load_avtolavki_chunks():
    """Отдельный chunk с агрегацией автолавок по городам-базам.

    В data/stores/all_stores.json у каждой автолавки address = адрес депо
    (откуда выезжает), не маршрут. Реальные маршруты Евроторгом не
    публикуются — мы можем дать список **городов-баз** с количеством
    автолавок и инструкцию обращения на горячую линию.

    Закрывает запросы тестера 28.04 / 01.05 «Автолавки Евроопт» — раньше
    бот отвечал общими словами, теперь даст конкретику.
    """
    fp = DATA / "stores/all_stores.json"
    if not fp.exists():
        return 0
    try:
        data = json.load(open(fp, encoding="utf-8"))
    except Exception:
        return 0

    # Собираем автолавки
    avto = [
        s for s in data
        if (s.get("format") or "").lower() in ("автолавка", "avtolavka")
        or "автолавк" in ((s.get("format") or "") + (s.get("brand") or "")).lower()
    ]
    if not avto:
        return 0

    # Группируем по городу-базе
    from collections import Counter
    cities = Counter(s.get("city", "—") for s in avto)
    city_lines = [
        f"• {city} — {count} автолавок{'а' if count == 1 else ('и' if 2 <= count <= 4 else '')}"
        for city, count in cities.most_common()
    ]

    text = (
        f"Автолавки сети «Евроопт» — мобильные магазины, обслуживающие "
        f"сельские населённые пункты Беларуси по расписанию.\n\n"
        f"Всего в сети **{len(avto)} автолавок**, базы расположены в {len(cities)} городах:\n"
        + "\n".join(city_lines) + "\n\n"
        f"Каждая автолавка обслуживает **несколько сёл** по расписанию — "
        f"маршруты и расписание автолавок **не публикуются на сайте**, "
        f"их можно уточнить на горячей линии Евроторга: **+375 44 788 88 80**.\n\n"
        f"В автолавках продаются: продукты питания первой необходимости, "
        f"бытовая химия, средства гигиены, корма для животных. Ассортимент "
        f"отличается от стационарных магазинов и может меняться от выезда к выезду.\n\n"
        f"Подробности и общая информация о сети — на https://evroopt.by/.\n\n"
        f"Ключевые слова: автолавки, автолавка, мобильные магазины, "
        f"сельская торговля, сёла, деревни, выездная торговля, маршрут автолавки, "
        f"расписание автолавки, когда приедет автолавка, есть ли автолавка"
    )
    add(
        "avtolavki_summary",
        text,
        "stores",
        source="all_stores.json",
        format="Автолавка",
    )
    return 1


def load_udacha():
    fp = DATA / "udacha/igra.evroopt УВП.xlsx"
    wb = openpyxl.load_workbook(fp)
    n = 0

    # FAQ-разделы УВП
    ws = wb["Основное разделы"]
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
        topic, desc, url = row[:3]
        if not topic:
            continue
        text = (
            f"«Удача в придачу!» (акция Евроторг): {topic}\n\n"
            f"{desc or ''}\n\n"
            f"Ссылка: {url or 'https://igra.evroopt.by/'}\n\n"
            f"Ключевые слова: Удача в придачу, УВП, игра, призы, туры, бонусы, Евроопт"
        )
        add(f"udacha_topic_{i}", text, "udacha", source="udacha_xlsx")
        n += 1

    # Товары удачи по сетям
    for sheet_name in wb.sheetnames:
        if not sheet_name.startswith("Товары удачи"):
            continue
        network = sheet_name.replace("Товары удачи ", "")
        ws = wb[sheet_name]
        items = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            code = row[0]
            name = row[1] if len(row) > 1 else None
            if code and name:
                items.append(f"{name}")
        if not items:
            continue
        # Один большой чанк на сеть (товаров много, заголовков мало)
        chunk_size = 20
        for chunk_i in range(0, len(items), chunk_size):
            chunk = items[chunk_i:chunk_i + chunk_size]
            text = (
                f"Товары «Удачи в придачу!» — {network} ({chunk_i+1}–{chunk_i+len(chunk)} из {len(items)}):\n\n"
                + "\n".join(f"• {x}" for x in chunk)
                + "\n\nЭти товары участвуют в розыгрыше «Удача в придачу!».\n"
                + f"Ключевые слова: Удача в придачу, товары удачи, {network}"
            )
            add(f"udacha_items_{network}_{chunk_i}", text, "udacha_items",
                source="udacha_xlsx", network=network)
            n += 1
    return n


# ============ Магазины с разметкой Евроопт / Хит ============

def load_stores_with_brand():
    """Новый файл с полем 'Тип склада(Общий)' → Евроопт / Хит."""
    fp = DATA / "stores_new/Список ТО ЕТ Хит с форматами.xlsx"
    wb = openpyxl.load_workbook(fp)
    ws = wb.active
    n = 0
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
        sklad, name, fmt, type_, brand = row[:5]
        if not name:
            continue
        m = re.search(r"Магазин\s+(.*?)(?::\(|$)", str(name))
        addr = m.group(1).strip() if m else str(name)
        if not addr:
            continue
        text = (
            f"Магазин {brand or 'сети Евроторг'}\n"
            f"Адрес: {addr}\n"
            f"Формат: {fmt or '—'}\n"
            f"Тип: {type_ or '—'}\n\n"
            f"Ключевые слова: магазин {brand or 'Евроторг'}, адрес"
        )
        add(f"store_branded_{i}", text, "store",
            source="stores_branded", brand=str(brand) if brand else "Евроторг")
        n += 1
    return n


def load_stores_minsk_full():
    """Магазины Минск с часами работы (из старого json) + mapping по районам/метро."""
    fp = DATA / "stores/euroopt.json"
    if not fp.exists():
        return 0
    items = json.load(open(fp, encoding="utf-8"))
    n = 0
    for i, st in enumerate(items):
        name = st.get("name", "Евроопт")
        addr = st.get("address", "")
        city = st.get("city", "Минск")
        hrs = st.get("hours", "")
        text = (
            f"Магазин Евроопт в {city}\n"
            f"Название: {name}\n"
            f"Адрес: {city}, {addr}\n"
            f"Часы работы: {hrs or 'уточните на evroopt.by'}\n\n"
            f"Ключевые слова: адрес Евроопт {city}, время работы, график"
        )
        add(f"store_minsk_hours_{i}", text, "store",
            source="minsk_json", brand="Евроопт", city=city)
        n += 1

    # ➕ Сводный чанк с районами Минска и станциями метро
    districts_text = (
        "Магазины «Евроопт» в Минске по районам и станциям метро:\n\n"
        "• Центр / пр. Независимости: ул. Независимости 48, Калиновского 23\n"
        "• Район Михалово / ст.м. Михалово: ул. Голубева 14, ул. Ауэзова 7/1\n"
        "• Район Академия наук / ст.м. Академия наук: ул. Независимости 74-98\n"
        "• Район Каменная горка / ст.м. Каменная Горка: ул. Притыцкого, Ольшевского\n"
        "• Район Уручье / ст.м. Уручье: ул. Голодеда 7/2, Независимости\n"
        "• Район Серебрянка / ст.м. Автозаводская: ул. Прушинских 2\n"
        "• Район Юго-Запад / ст.м. Петровщина-Малиновка: ул. Есенина 141, ул. Алибегова 13/1\n"
        "• Район Ленинский / ст.м. Пушкинская: ул. Багратиона 55Б\n"
        "• Район Фрунзенский / ст.м. Фрунзенская: ул. Грицевца 1-66\n"
        "• Ложинская / ст.м. Борисовский тракт: ул. Ложинская 22\n\n"
        "Полный список и часы работы — на сайте https://evroopt.by/ в разделе «Магазины». "
        "Адреса магазинов «Хит Дискаунтер» — на hitdiscount.by, «Грошык» — на groshyk.by.\n\n"
        "Ключевые слова: магазин Минск район, около метро, станция метро, где ближайший, "
        "Михалово, Академия наук, Каменная горка, Уручье, Серебрянка, Юго-Запад, Петровщина, "
        "Малиновка, Пушкинская, Фрунзенская, Борисовский тракт, центр Минска, рядом со мной"
    )
    add("stores_minsk_districts", districts_text, "store",
        source="districts_map", brand="Евроопт", city="Минск")
    n += 1
    return n


def load_all_stores_v2():
    """Загрузка нового справочника 1040 магазинов с brand/format/city разметкой.

    Источник: data/stores/all_stores.json (генерируется scripts/parse_stores_xlsx.py).
    Формат: {brand, format, city, address, raw_name}.
    Закрывает 24.04 P1 (смешение Евроопт/Хит) — каждая запись имеет
    metadata.brand для жёсткой фильтрации в RAG.search.
    """
    fp = DATA / "stores/all_stores.json"
    if not fp.exists():
        return 0
    items = json.load(open(fp, encoding="utf-8"))
    n = 0
    for s in items:
        brand = s.get("brand") or "Евроопт"
        fmt = s.get("format") or "Магазин"
        city = s.get("city") or ""
        address = s.get("address") or ""
        sid = s.get("id", n)
        text = (
            f"Магазин сети {brand}\n"
            f"Формат: {fmt}\n"
            f"Город: {city}\n"
            f"Адрес: {address}\n\n"
            f"Ключевые слова: {brand} {city}, {fmt} в {city}, "
            f"магазин в {city}, адрес {brand} {city}"
        )
        add(
            f"store_v2_{sid}",
            text,
            "store",
            source="all_stores",
            brand=brand,
            format=fmt,
            city=city,
        )
        n += 1
    return n


def load_general_faq_additions():
    """Дополнительные FAQ-записи, закрывающие повторяющиеся вопросы тестеров."""
    additions = [
        ("past_promo",
         "О прошлых и архивных акциях",
         "Архив прошедших акций в базе не хранится. Информация о действующих акциях "
         "актуализируется каждые 2 недели. Для истории промо — воспользуйтесь "
         "разделом https://evroopt.by/news/ или https://hitdiscount.by/.",
         "прошлые акции, архив акций, что было, прошлая неделя, завершённые промо"),

        ("future_promo",
         "О будущих акциях и праздничных предложениях",
         "Информация о предстоящих акциях (на Пасху, Новый Год, 8 марта, 14 февраля, "
         "День Победы, другие праздники) появляется за 1–2 недели до старта. "
         "Следите за анонсами на evroopt.by/news/ и подписывайтесь на программу Еплюс.",
         "Пасха акции, новогодние скидки, 8 марта, 14 февраля, праздничные акции, "
         "предстоящие акции, будут ли скидки, анонсы"),

        ("how_find_store",
         "Как найти ближайший магазин",
         "Чтобы найти ближайший магазин с геолокацией, воспользуйтесь картой "
         "на сайтах evroopt.by, hitdiscount.by, groshyk.by — там можно искать "
         "по адресу, районам и станциям метро. Приложение «Еплюс» тоже показывает "
         "магазины рядом с вашим местоположением.",
         "ближайший магазин, рядом со мной, найти магазин, карта магазинов, "
         "где купить, локация, геолокация"),

        ("pay_methods",
         "Способы оплаты в Евроопт, Грошык, Хит",
         "Во всех магазинах сети Евроторг («Евроопт», «Грошык», «Хит Дискаунтер») "
         "принимаются: наличные, банковские карты Visa/Mastercard/Белкарт, оплата "
         "по QR-коду (ЕРИП, Оплати), Apple Pay, Google Pay, Samsung Pay. "
         "Бонусами «Еплюс» можно оплатить до 99% стоимости товара.",
         "как оплатить, способы оплаты, Apple Pay, Google Pay, QR, ЕРИП, белкарт, "
         "карта, наличные, оплата бонусами"),

        ("delivery_summary",
         "Доставка продуктов",
         "Доставка продуктов оформляется через сайт Евроопт https://evroopt.by/ "
         "в разделе заказа онлайн. Сумма минимального заказа, зоны и время доставки "
         "уточняются при оформлении — зависят от адреса и магазина. "
         "Оплата: картой онлайн либо при получении наличными/картой.",
         "доставка продуктов, доставка Евроопт, заказ на дом, оформить доставку, "
         "онлайн-заказ, минимальный заказ, зоны доставки"),

        ("eplus_levels",
         "Привилегии по карте Еплюс",
         "Карта «Еплюс» — бесплатная (виртуальная) или 99 копеек (пластик). "
         "Накопленные бонусы (1 бонус = 1 копейка) можно тратить на оплату до 99% "
         "стоимости товара, кроме алкоголя и табака. Бонусы не начисляются за "
         "алкоголь, табак, подарочные сертификаты, лотерейные билеты, SIM-карты. "
         "Срок жизни стандартных бонусов — 365 дней. Полный FAQ на eplus.by.",
         "Еплюс привилегии, бонусы Еплюс, как копить, как тратить, срок бонусов, "
         "что нельзя купить бонусами, что не начисляется"),
    ]
    n = 0
    for key, title, body, keywords in additions:
        text = (
            f"{title}\n\n{body}\n\n"
            f"Ключевые слова: {keywords}"
        )
        add(f"faq_add_{key}", text, "faq", source="manual_additions")
        n += 1
    return n


# ============ Рецепты ============

def load_recipes_docx():
    fp = DATA / "recipes/recipes_alexey.docx"
    if not fp.exists():
        return 0
    d = Document(fp)

    # Парсим по блокам "Название рецепта — X порций — калорийность — Ингредиенты ... — Шаги"
    # Ищем по маркерам "Порции:", "Ингредиенты:", "Приготовление:", "Шаг"
    current = None
    recipes = []
    for p in d.paragraphs:
        text = p.text.strip()
        if not text:
            if current and len(current.get("lines", [])) > 3:
                # Пустая строка — возможно разделитель
                pass
            continue
        # Начало нового рецепта — короткая строка без маркеров
        if (
            len(text) < 60 and
            not text.startswith(("Шаг", "•", "-", "–", "Ингредиенты", "Приготовление", "Порции", "Калорийность", "Время"))
            and not re.match(r"^\d+[.\)]", text)
            and ":" not in text
            and len(text.split()) <= 6
            and any(c.isupper() for c in text[:3])
        ):
            if current and len(current["lines"]) > 3:
                recipes.append(current)
            current = {"name": text, "lines": []}
            continue
        if current:
            current["lines"].append(text)
    if current and len(current["lines"]) > 3:
        recipes.append(current)

    n = 0
    for i, r in enumerate(recipes):
        name = r["name"]
        if len(name) < 4:
            continue
        body = "\n".join(r["lines"])
        text = (
            f"Рецепт: {name}\n\n{body}\n\n"
            f"Ключевые слова: рецепт, {name}, как приготовить, ингредиенты"
        )
        add(f"recipe_v2_{i}", text, "recipe", source="recipes_alexey", name=name)
        n += 1
    return n


def load_recipes_json():
    n = 0
    for fn in ["sample.json", "youtube.json"]:
        fp = DATA / "recipes" / fn
        if not fp.exists():
            continue
        items = json.load(open(fp, encoding="utf-8"))
        if isinstance(items, dict):
            items = [items]
        for i, r in enumerate(items):
            name = r.get("name") or r.get("title", "Рецепт")
            ingr = r.get("ingredients", [])
            steps = r.get("steps", [])
            url = r.get("url", "")
            text = f"Рецепт: {name}\n"
            if ingr:
                text += "Ингредиенты:\n" + "\n".join(f"- {i}" for i in ingr) + "\n"
            if steps:
                text += "Приготовление:\n" + "\n".join(f"{j+1}. {s}" for j, s in enumerate(steps))
            if url:
                text += f"\nВидео: {url}"
            text += f"\n\nКлючевые слова: рецепт, {name}"
            add(f"recipe_{fn}_{i}", text, "recipe", source=fn, name=name)
            n += 1
    return n


def main():
    print("=== Парсинг источников ===")
    counts = {
        "FAQ Еплюс": load_faq_eplus(),
        "FAQ JSON (с заменой e-доставки→evroopt)": load_faq_json(),
        "🆕 FAQ additions (районы, прошлые/будущие акции, оплата)": load_general_faq_additions(),
        "Brand links (evroopt/groshyk/hitdiscount)": load_brand_links(),
        "🆕 Акции Евроопт (Красная цена, Цены вниз, Бонусы, Грильфест)": load_promotions_evroopt(),
        "🆕 Акции Хит Дискаунтер": load_promotions_hit(),
        "🆕 Текущая листовка Евроопт (snapshot от заказчика)": load_listovka_current(),
        "🆕 Автолавки — агрегатный chunk по городам": load_avtolavki_chunks(),
        "🆕 Удача в придачу — текущий тур (cron-парсер главной)": load_udacha_current_tur(),
        "🆕 Удача в придачу (FAQ + товары удачи)": load_udacha(),
        "🆕 Магазины с брендом (Евроопт/Хит)": load_stores_with_brand(),
        "🆕🆕 v2: 1040 магазинов с brand/format/city": load_all_stores_v2(),
        "Магазины Минск с часами": load_stores_minsk_full(),
        "Рецепты DOCX": load_recipes_docx(),
        "Рецепты JSON": load_recipes_json(),
    }
    for k, v in counts.items():
        print(f"  {k}: {v}")
    print(f"  ИТОГО: {len(docs)} документов")

    print("\n=== Загрузка в RAG (новая коллекция v3) ===")
    rag = RAGEngine()
    try:
        rag.client.delete_collection("euroopt_knowledge_v3")
        print("  старая коллекция удалена")
    except Exception:
        pass
    rag = RAGEngine()
    for i in range(0, len(docs), 50):
        rag.add_documents(docs[i:i+50])
    print(f"  загружено: {rag.collection.count()}")

    print("\n=== Тест поиска ===")
    for q in [
        "Что такое Красная цена?", "какие акции в Хит?", "товары удачи",
        "Цены вниз мясо", "магазин Хит в Минске", "призы в 214 туре",
        "Какая скидка на капусту?", "какие бонус-товары?",
    ]:
        results = rag.search(q, n_results=2)
        print(f"\nQ: {q} ({len(results)})")
        for r in results[:2]:
            print(f"  score={r.get('score',0):.3f} | {(r.get('text') or '')[:120]}")


if __name__ == "__main__":
    main()
